#!/usr/bin/env python3
"""Build DOCX templates with SmartForm placeholders from collected public forms."""

from pathlib import Path
from subprocess import run
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "template"
OUTPUT = SOURCE / "renderable"

SPECS = {
    "HT_KHAI_SINH": {
        "source": "1._To_khai_dang_ky_khai_sinh_1605163058.doc",
        "paragraphs": {
            6: "Kính gửi: {{coQuanTiepNhan}}",
            8: "Họ, chữ đệm, tên người yêu cầu: {{nguoiYeuCau.hoTen}}",
            9: "Ngày, tháng, năm sinh: {{nguoiYeuCau.ngaySinh}}",
            10: "Nơi cư trú: {{nguoiYeuCau.noiCuTru}}",
            12: "Giấy tờ tùy thân: CCCD số {{nguoiYeuCau.soCCCD}}",
            14: "Quan hệ với người được khai sinh: {{nguoiYeuCau.quanHe}}",
            16: "Họ, chữ đệm, tên trẻ: {{tre.hoTen}}",
            17: "Ngày, tháng, năm sinh: {{tre.ngaySinh}}",
            19: "Giới tính: {{tre.gioiTinh}}    Dân tộc: {{tre.danToc}}    Quốc tịch: {{tre.quocTich}}",
            20: "Nơi sinh: {{tre.noiSinh}}",
            22: "Quê quán: {{tre.queQuan}}",
            23: "Họ, chữ đệm, tên người mẹ: {{me.hoTen}}",
            27: "Giấy tờ tùy thân của mẹ: {{me.soCCCD}}",
            28: "Họ, chữ đệm, tên người cha: {{cha.hoTen}}",
            33: "Thông tin Giấy chứng nhận kết hôn: Số {{honNhan.soGiayTo}}",
        },
    },
    "HT_TRICH_LUC_KHAI_SINH": {
        "source": "mau-to-khai-cap-ban-sao-trich-luc-ho-tich.doc",
        "paragraphs": {
            6: "Kính gửi: {{coQuanTiepNhan}}",
            8: "Họ, chữ đệm, tên người yêu cầu: {{nguoiYeuCau.hoTen}}",
            9: "Nơi cư trú: {{nguoiYeuCau.noiCuTru}}",
            11: "Giấy tờ tùy thân: CCCD số {{nguoiYeuCau.soCCCD}}",
            13: "Quan hệ với người được cấp bản sao: {{nguoiYeuCau.quanHe}}",
            14: "Đề nghị cấp bản sao trích lục khai sinh cho người có tên dưới đây:",
            15: "Họ, chữ đệm, tên: {{hoTich.hoTen}}",
            16: "Ngày, tháng, năm sinh: {{hoTich.ngaySinh}}",
            17: "Giới tính: {{hoTich.gioiTinh}}",
            23: "Đã đăng ký tại: {{hoTich.noiDangKy}}",
            25: "Ngày đăng ký: {{hoTich.ngayDangKy}}    Số: {{hoTich.so}}    Quyển số: {{hoTich.quyenSo}}",
            26: "Số lượng bản sao yêu cầu cấp: {{yeuCau.soLuong}} bản.",
        },
    },
    "HT_DK_LAI_KHAI_SINH": {
        "source": "1._To_khai_dang_ky_khai_sinh_1605163058.doc",
        "paragraphs": {
            6: "Kính gửi: {{coQuanTiepNhan}}",
            8: "Họ, chữ đệm, tên người yêu cầu: {{nguoiYeuCau.hoTen}}",
            9: "Ngày, tháng, năm sinh: {{nguoiYeuCau.ngaySinh}}",
            10: "Nơi cư trú: {{nguoiYeuCau.noiCuTru}}",
            12: "Giấy tờ tùy thân: CCCD số {{nguoiYeuCau.soCCCD}}",
            14: "Quan hệ với người được khai sinh: {{nguoiYeuCau.quanHe}}",
            16: "Họ, chữ đệm, tên: {{hoTich.hoTen}}",
            17: "Ngày, tháng, năm sinh: {{hoTich.ngaySinh}}",
            19: "Giới tính: {{hoTich.gioiTinh}}",
            20: "Nơi đã đăng ký khai sinh trước đây: {{hoTich.noiDangKyCu}}",
            23: "Họ, chữ đệm, tên người mẹ: {{me.hoTen}}",
            28: "Họ, chữ đệm, tên người cha: {{cha.hoTen}}",
            33: "Lý do đăng ký lại: {{lyDo.dangKyLai}}",
        },
    },
    "HT_XNTTHN": {
        "source": "mau-to-khai-cap-giay-xac-nhan-tinh-trang-hon-nhan.docx",
        "paragraphs": {
            4: "Kính gửi: {{coQuanTiepNhan}}",
            6: "Họ, chữ đệm, tên người yêu cầu: {{nguoiYeuCau.hoTen}}",
            7: "Nơi cư trú: {{nguoiYeuCau.noiCuTru}}",
            9: "Giấy tờ tùy thân: CCCD số {{nguoiYeuCau.soCCCD}}",
            11: "Quan hệ với người được cấp Giấy xác nhận: {{nguoiYeuCau.quanHe}}",
            13: "Họ, chữ đệm, tên: {{nguoiDuocXacNhan.hoTen}}",
            14: "Ngày, tháng, năm sinh: {{nguoiDuocXacNhan.ngaySinh}}",
            15: "Giới tính: {{nguoiDuocXacNhan.gioiTinh}}",
            16: "Nơi cư trú: {{nguoiDuocXacNhan.noiCuTru}}",
            18: "Giấy tờ tùy thân: CCCD số {{nguoiYeuCau.soCCCD}}",
            20: "Tình trạng hôn nhân: {{honNhan.tinhTrang}}",
            25: "Mục đích sử dụng Giấy xác nhận tình trạng hôn nhân: {{honNhan.mucDich}}",
        },
    },
}


# Biểu mẫu hộ kinh doanh theo Phụ lục II Thông tư 68/2025/TT-BTC.
# Placeholder {{...}} khớp formFields trong mvp-procedures.ts; các trường lấy từ
# giấy chứng nhận cũ tuân theo template/document-types/HO_KINH_DOANH.json.
BUILT_SPECS = {
    "HKD_THANH_LAP": {
        "title": "GIẤY ĐỀ NGHỊ ĐĂNG KÝ HỘ KINH DOANH",
        "subtitle": "(Mẫu số 1, Phụ lục II, Thông tư 68/2025/TT-BTC)",
        "lines": [
            "Kính gửi: {{coQuanTiepNhan}}",
            "",
            "Tôi là (ghi họ tên bằng chữ in hoa): {{chuHo.hoTen}}",
            "Sinh ngày: {{chuHo.ngaySinh}}",
            "Số định danh cá nhân: {{chuHo.soCCCD}}",
            "Nơi thường trú: {{chuHo.diaChiThuongTru}}",
            "",
            "Đăng ký hộ kinh doanh do tôi là chủ hộ với các nội dung sau:",
            "1. Tên hộ kinh doanh viết bằng tiếng Việt (ghi bằng chữ in hoa): {{hoKinhDoanh.tenHoKinhDoanh}}",
            "2. Trụ sở của hộ kinh doanh: {{hoKinhDoanh.diaChiKinhDoanh}}",
            "3. Ngành, nghề kinh doanh: {{hoKinhDoanh.nganhNghe}}",
            "4. Vốn kinh doanh (bằng số, bằng chữ, VNĐ): {{hoKinhDoanh.vonKinhDoanh}}",
            "5. Chủ thể thành lập hộ kinh doanh: Cá nhân / Các thành viên hộ gia đình",
            "",
            "Tôi cam kết thông tin kê khai là hợp pháp, chính xác và trung thực.",
        ],
        "signature": "CHỦ HỘ KINH DOANH\n(Ký và ghi họ tên)",
    },
    "HKD_THAY_DOI": {
        "title": "GIẤY ĐỀ NGHỊ\nĐĂNG KÝ THAY ĐỔI NỘI DUNG ĐĂNG KÝ HỘ KINH DOANH",
        "subtitle": "(Mẫu số 2, Phụ lục II, Thông tư 68/2025/TT-BTC)",
        "lines": [
            "Kính gửi: {{coQuanTiepNhan}}",
            "",
            "Tên hộ kinh doanh (ghi bằng chữ in hoa): {{hoKinhDoanh.tenHoKinhDoanh}}",
            "Mã số hộ kinh doanh/Mã số đăng ký hộ kinh doanh/Mã số thuế: {{hoKinhDoanh.maSo}}",
            "Trụ sở hiện tại: {{hoKinhDoanh.diaChiHienTai}}",
            "",
            "Hộ kinh doanh đăng ký thay đổi nội dung như sau:",
            "- Nội dung thay đổi: {{thayDoi.noiDung}}",
            "- Thông tin sau khi thay đổi: {{thayDoi.giaTriMoi}}",
            "",
            "Chủ hộ kinh doanh: {{chuHo.hoTen}}",
            "Số định danh cá nhân: {{chuHo.soCCCD}}",
            "Hộ kinh doanh cam kết nội dung đề nghị là hợp pháp, chính xác và trung thực.",
        ],
        "signature": "CHỦ HỘ KINH DOANH\n(Ký và ghi họ tên)",
    },
    "HKD_CAP_LAI": {
        "title": "GIẤY ĐỀ NGHỊ CẤP LẠI GIẤY CHỨNG NHẬN ĐĂNG KÝ HỘ KINH DOANH",
        "subtitle": "(Mẫu số 5, Phụ lục II, Thông tư 68/2025/TT-BTC)",
        "lines": [
            "Kính gửi: {{coQuanTiepNhan}}",
            "",
            "Tên hộ kinh doanh (ghi bằng chữ in hoa): {{hoKinhDoanh.tenHoKinhDoanh}}",
            "Mã số hộ kinh doanh: {{hoKinhDoanh.maSo}}",
            "Đề nghị được cấp lại Giấy chứng nhận đăng ký hộ kinh doanh.",
            "Lý do đề nghị cấp lại: {{capLai.lyDo}}",
            "",
            "Hộ kinh doanh cam kết nội dung đề nghị là hợp pháp, chính xác và trung thực.",
        ],
        "signature": "CHỦ HỘ KINH DOANH\n(Ký và ghi họ tên)",
    },
    "HKD_CHAM_DUT": {
        "title": "THÔNG BÁO VỀ VIỆC CHẤM DỨT HOẠT ĐỘNG HỘ KINH DOANH",
        "subtitle": "(Mẫu số 4, Phụ lục II, Thông tư 68/2025/TT-BTC)",
        "lines": [
            "Kính gửi: {{coQuanTiepNhan}}",
            "",
            "Tên hộ kinh doanh (ghi bằng chữ in hoa): {{hoKinhDoanh.tenHoKinhDoanh}}",
            "Mã số hộ kinh doanh: {{hoKinhDoanh.maSo}}",
            "",
            "Hộ kinh doanh cam kết đã thanh toán đầy đủ các khoản nợ, gồm cả nợ thuế và nghĩa vụ tài chính khác, hoặc đã đạt được thỏa thuận với các chủ nợ về việc thực hiện nghĩa vụ tài chính trước khi nộp hồ sơ.",
            "",
            "Hộ kinh doanh cam kết nội dung thông báo là hợp pháp, chính xác và trung thực.",
        ],
        "signature": "CHỦ HỘ KINH DOANH/NGƯỜI THỪA KẾ\n(Ký và ghi họ tên)",
    },
}


def build_from_scratch(spec: dict) -> Document:
    document = Document()
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    run.bold = True
    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Độc lập - Tự do - Hạnh phúc").bold = True
    document.add_paragraph("")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(spec["title"]).bold = True
    if spec.get("subtitle"):
        note = document.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note.add_run(spec["subtitle"]).italic = True
    document.add_paragraph("")

    for line in spec["lines"]:
        document.add_paragraph(line)

    document.add_paragraph("")
    signature = document.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature.add_run(spec.get("signature", "Người yêu cầu\n(Ký, ghi rõ họ tên)")).italic = True
    return document


def as_docx(source: Path, temp: Path) -> Path:
    if source.suffix.lower() == ".docx":
        return source
    run(
        ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", str(temp), str(source)],
        check=True,
        capture_output=True,
    )
    return temp / f"{source.stem}.docx"


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    with TemporaryDirectory() as directory:
        temp = Path(directory)
        for key, spec in SPECS.items():
            document = Document(as_docx(SOURCE / spec["source"], temp))
            for index, text in spec["paragraphs"].items():
                document.paragraphs[index].text = text
            for section in document.sections:
                section.page_width = Mm(210)
                section.page_height = Mm(297)
            destination = OUTPUT / f"{key}.docx"
            document.save(destination)
            print(destination.relative_to(ROOT))

    for key, spec in BUILT_SPECS.items():
        document = build_from_scratch(spec)
        for section in document.sections:
            section.page_width = Mm(210)
            section.page_height = Mm(297)
        destination = OUTPUT / f"{key}.docx"
        document.save(destination)
        print(destination.relative_to(ROOT))


if __name__ == "__main__":
    main()
