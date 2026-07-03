#!/usr/bin/env python3
"""
Script đầy đủ để thêm tất cả placeholder vào HKD_THAY_DOI template
"""

from docx import Document

def add_all_placeholders():
    input_file = "/home/dangkien/1st-Main/hackaithon/GovTrust-AI/template/renderable/HKD_THAY_DOI.docx"
    output_file = "/home/dangkien/1st-Main/hackaithon/GovTrust-AI/template/renderable/HKD_THAY_DOI.docx"

    doc = Document(input_file)

    # Dictionary đầy đủ tất cả các thay thế
    replacements = {
        # Header
        'TÊN HỘ KINH DOANH\n-------': '{{hoKinhDoanh.tenHoKinhDoanh}}',

        # Thông tin hộ kinh doanh (chưa có placeholder)
        'Điện thoại (nếu có): ..................................................................  Fax (nếu có): ..............':
            'Điện thoại: {{hoKinhDoanh.dienThoai}}  Fax: {{hoKinhDoanh.fax}}',
        'Email (nếu có): ........................................................................  Website (nếu có): .......':
            'Email: {{hoKinhDoanh.email}}  Website: {{hoKinhDoanh.website}}',

        # Phần 1: Chủ hộ CŨ (còn thiếu nhiều trường)
        'Họ và tên (ghi bằng chữ in hoa): ...............................................  Giới tính: ....................':
            'Họ và tên: {{chuHoCu.hoTen}}  Giới tính: {{chuHoCu.gioiTinh}}',
        'Sinh ngày: ........./....../........ Dân tộc: ......  Quốc tịch: ....................':
            'Sinh ngày: {{chuHoCu.ngaySinh}} Dân tộc: {{chuHoCu.danToc}}  Quốc tịch: {{chuHoCu.quocTich}}',
        'Điện thoại (nếu có): ................................................  Email (nếu có): ..........................':
            'Điện thoại: {{chuHoCu.dienThoai}}  Email: {{chuHoCu.email}}',

        # Phần 2: Chủ hộ MỚI (còn thiếu nhiều trường)
        'Họ và tên (ghi bằng chữ in hoa): .....................................................  Giới tính: ……….':
            'Họ và tên: {{chuHoMoi.hoTen}}  Giới tính: {{chuHoMoi.gioiTinh}}',
        'Sinh ngày: ................ /............... /........... Dân tộc: ........................  Quốc tịch: ............':
            'Sinh ngày: {{chuHoMoi.ngaySinh}} Dân tộc: {{chuHoMoi.danToc}}  Quốc tịch: {{chuHoMoi.quocTich}}',
        'Số giấy tờ pháp lý của cá nhân: {{chuHoCu.soCCCD}}':
            'Số giấy tờ pháp lý của cá nhân: {{chuHoMoi.soCCCD}}',
        'Ngày cấp: ..../..../.... Nơi cấp: ..................................................................................':
            'Ngày cấp: {{chuHoMoi.ngayCap}} Nơi cấp: {{chuHoMoi.noiCap}}',
        'Có giá trị đến ngày: {{chuHoCu.hanSuDung}}':
            'Có giá trị đến ngày: {{chuHoMoi.hanSuDung}}',
        'Điện thoại (nếu có): .......................................................  Email (nếu có): ......................':
            'Điện thoại: {{chuHoMoi.dienThoai}}  Email: {{chuHoMoi.email}}',
    }

    replaced_count = 0

    # Process paragraphs
    for para in doc.paragraphs:
        original = para.text
        modified = original

        for old, new in replacements.items():
            if old in modified:
                modified = modified.replace(old, new)

        if modified != original:
            para.clear()
            para.add_run(modified)
            replaced_count += 1
            print(f"✓ [{replaced_count}] Replaced in paragraph")

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original = para.text
                    modified = original

                    for old, new in replacements.items():
                        if old in modified:
                            modified = modified.replace(old, new)

                    if modified != original:
                        para.clear()
                        para.add_run(modified)
                        replaced_count += 1
                        print(f"✓ [{replaced_count}] Replaced in table")

    doc.save(output_file)
    print(f"\n✅ Total replacements: {replaced_count}")
    print(f"✅ Saved to: {output_file}")

if __name__ == "__main__":
    print("🔧 Adding ALL placeholders to HKD_THAY_DOI.docx...\n")
    add_all_placeholders()
