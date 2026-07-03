#!/usr/bin/env python3
"""
Script tự động thêm placeholder vào template HKD_THAY_DOI.docx
Thay thế các dấu chấm/gạch ngang bằng {{placeholder}}
"""

from docx import Document
import re
import sys

def add_placeholders_to_hkd_template(input_file, output_file):
    """
    Thêm placeholder vào template HKD_THAY_DOI
    """
    doc = Document(input_file)

    # Mapping: simple text replacements (exact match)
    # Sử dụng simple string replace thay vì regex để tránh phức tạp
    replacements = {
        # Header
        'TÊN HỘ KINH DOANH-------': '{{hoKinhDoanh.tenHoKinhDoanh}}',
        'Số: .................': 'Số: {{soVanBan}}',
        '......, ngày...... tháng...... năm......': '{{ngayThang}}',

        # Thông tin hộ kinh doanh
        'Tên hộ kinh doanh (ghi bằng chữ in hoa): ......................................................................': 'Tên hộ kinh doanh: {{hoKinhDoanh.tenHoKinhDoanh}}',
        'Mã số hộ kinh doanh/Mã số thuế: ................................................................................': 'Mã số hộ kinh doanh/Mã số thuế: {{hoKinhDoanh.maSo}}',
        'Mã số đăng ký hộ kinh doanh: .....................................................................................': 'Mã số đăng ký hộ kinh doanh: {{hoKinhDoanh.maSoDangKy}}',
        'Điện thoại (nếu có): ..................................................................  Fax (nếu có): ..............': 'Điện thoại: {{hoKinhDoanh.dienThoai}}  Fax: {{hoKinhDoanh.fax}}',
        'Email (nếu có): ........................................................................  Website (nếu có): .......': 'Email: {{hoKinhDoanh.email}}  Website: {{hoKinhDoanh.website}}',

        # Chủ hộ cũ
        'Họ và tên (ghi bằng chữ in hoa): ...............................................  Giới tính: ....................': 'Họ và tên: {{chuHoCu.hoTen}}  Giới tính: {{chuHoCu.gioiTinh}}',
        'Sinh ngày: ........./....../........ Dân tộc: ......  Quốc tịch: ....................': 'Sinh ngày: {{chuHoCu.ngaySinh}} Dân tộc: {{chuHoCu.danToc}}  Quốc tịch: {{chuHoCu.quocTich}}',
        'Số giấy tờ pháp lý của cá nhân: ............................................................................': 'Số giấy tờ pháp lý của cá nhân: {{chuHoCu.soCCCD}}',
        'Ngày cấp: ..../..../.... Nơi cấp: ...................................................................................': 'Ngày cấp: {{chuHoCu.ngayCap}} Nơi cấp: {{chuHoCu.noiCap}}',
        'Có giá trị đến ngày (nếu có): .../.../...': 'Có giá trị đến ngày: {{chuHoCu.hanSuDung}}',
        'Số nhà, ngách, hẻm, ngõ, đường phố/tổ/xóm/ấp/thôn: .................................................': 'Số nhà, ngách, hẻm, ngõ, đường phố/tổ/xóm/ấp/thôn: {{diaChiThuongTru.soNha}}',
        'Xã/Phường/Thị trấn: ....................................................................................................': 'Xã/Phường/Thị trấn: {{diaChiThuongTru.xaPhuong}}',
        'Quận/Huyện/Thị xã/Thành phố thuộc tỉnh: .....................................................................': 'Quận/Huyện/Thị xã/Thành phố thuộc tỉnh: {{diaChiThuongTru.quanHuyen}}',
        'Tỉnh/Thành phố: ...........................................................................................................': 'Tỉnh/Thành phố: {{diaChiThuongTru.tinhThanh}}',
        'Điện thoại (nếu có): ................................................  Email (nếu có): .............................': 'Điện thoại: {{dienThoai}}  Email: {{email}}',
    }

    # Xử lý từng paragraph
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        modified_text = original_text

        for old_text, new_text in replacements.items():
            if old_text in modified_text:
                modified_text = modified_text.replace(old_text, new_text)

        if modified_text != original_text:
            # Clear paragraph và thêm lại với text mới
            paragraph.clear()
            paragraph.add_run(modified_text)
            print(f"✓ Replaced in paragraph")

    # Xử lý tables (nếu có)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    modified_text = original_text

                    for old_text, new_text in replacements.items():
                        if old_text in modified_text:
                            modified_text = modified_text.replace(old_text, new_text)

                    if modified_text != original_text:
                        paragraph.clear()
                        paragraph.add_run(modified_text)

    # Save
    doc.save(output_file)
    print(f"\n✅ Saved to: {output_file}")

if __name__ == "__main__":
    input_file = "/home/dangkien/1st-Main/hackaithon/GovTrust-AI/template/renderable/HKD_THAY_DOI.docx"
    output_file = "/home/dangkien/1st-Main/hackaithon/GovTrust-AI/template/renderable/HKD_THAY_DOI_with_placeholders.docx"

    print("🔧 Adding placeholders to HKD_THAY_DOI.docx...")
    add_placeholders_to_hkd_template(input_file, output_file)
    print("\n✅ Done! Please review the output file and replace the original if it looks correct.")
    print(f"   cp {output_file} {input_file}")
