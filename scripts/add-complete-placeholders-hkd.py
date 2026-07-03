#!/usr/bin/env python3
"""
Script hoàn chỉnh để thêm TẤT CẢ placeholder vào HKD_THAY_DOI template
Dựa trên cấu trúc trong mau-thong-bao-thay-di-chu-ho-kinh-doanh.md
"""

from docx import Document
import re

def add_complete_placeholders():
    input_file = "/home/dangkien/1st-Main/hackaithon/GovTrust-AI/template/renderable/HKD_THAY_DOI.docx"
    output_file = "/home/dangkien/1st-Main/hackaithon/GovTrust-AI/template/renderable/HKD_THAY_DOI.docx"

    doc = Document(input_file)

    # Danh sách thay thế THEO THỨ TỰ (quan trọng để tránh nhầm lẫn)
    replacements = [
        # === HEADER ===
        ('TÊN HỘ KINH DOANH\n-------', '{{hoKinhDoanh.tenHoKinhDoanh}}'),

        # === THÔNG TIN HỘ KINH DOANH ===
        ('Tên hộ kinh doanh (ghi bằng chữ in hoa): ......................................................................',
         'Tên hộ kinh doanh: {{hoKinhDoanh.tenHoKinhDoanh}}'),
        ('Mã số hộ kinh doanh/Mã số thuế: ................................................................................',
         'Mã số hộ kinh doanh/Mã số thuế: {{hoKinhDoanh.maSo}}'),
        ('Mã số đăng ký hộ kinh doanh: .....................................................................................',
         'Mã số đăng ký hộ kinh doanh: {{hoKinhDoanh.maSoDangKy}}'),

        # === PHẦN 1: CHỦ HỘ CŨ ===
        # Họ tên và giới tính
        ('Họ và tên (ghi bằng chữ in hoa): ...............................................  Giới tính: ....................',
         'Họ và tên: {{chuHoCu.hoTen}}  Giới tính: {{chuHoCu.gioiTinh}}'),

        # Sinh ngày, dân tộc, quốc tịch
        ('Sinh ngày: ........./....../........ Dân tộc: ......  Quốc tịch: ....................',
         'Sinh ngày: {{chuHoCu.ngaySinh}} Dân tộc: {{chuHoCu.danToc}}  Quốc tịch: {{chuHoCu.quocTich}}'),

        # Số giấy tờ (đầu tiên - chủ hộ cũ)
        ('Số giấy tờ pháp lý của cá nhân: ............................................................................',
         'Số giấy tờ pháp lý của cá nhân: {{chuHoCu.soCCCD}}'),

        # Ngày cấp, nơi cấp (đầu tiên - chủ hộ cũ)
        ('Ngày cấp: ..../..../.... Nơi cấp: ...................................................................................',
         'Ngày cấp: {{chuHoCu.ngayCap}} Nơi cấp: {{chuHoCu.noiCap}}'),

        # Có giá trị đến ngày (đầu tiên - chủ hộ cũ)
        ('Có giá trị đến ngày (nếu có): .../.../...',
         'Có giá trị đến ngày: {{chuHoCu.hanSuDung}}'),

        # Địa chỉ thường trú chủ hộ cũ (xuất hiện 2 lần - lần 1)
        # Điện thoại email (đầu tiên - chủ hộ cũ)
        ('Điện thoại (nếu có): ................................................  Email (nếu có): ...........................',
         'Điện thoại: {{chuHoCu.dienThoai}}  Email: {{chuHoCu.email}}'),

        # === PHẦN 2: CHỦ HỘ MỚI ===
        # Họ tên và giới tính
        ('Họ và tên (ghi bằng chữ in hoa): .....................................................  Giới tính: ……….',
         'Họ và tên: {{chuHoMoi.hoTen}}  Giới tính: {{chuHoMoi.gioiTinh}}'),

        # Sinh ngày, dân tộc, quốc tịch
        ('Sinh ngày: ................ /............... /........... Dân tộc: ........................  Quốc tịch: ............',
         'Sinh ngày: {{chuHoMoi.ngaySinh}} Dân tộc: {{chuHoMoi.danToc}}  Quốc tịch: {{chuHoMoi.quocTich}}'),

        # Số giấy tờ (thứ hai - chủ hộ mới)
        ('Số giấy tờ pháp lý của cá nhân: ............................................................................',
         'Số giấy tờ pháp lý của cá nhân: {{chuHoMoi.soCCCD}}'),

        # Ngày cấp, nơi cấp (thứ hai - chủ hộ mới)
        ('Ngày cấp: ..../..../.... Nơi cấp: ..................................................................................',
         'Ngày cấp: {{chuHoMoi.ngayCap}} Nơi cấp: {{chuHoMoi.noiCap}}'),

        # Có giá trị đến ngày (thứ hai - chủ hộ mới)
        ('Có giá trị đến ngày (nếu có): .../.../...',
         'Có giá trị đến ngày: {{chuHoMoi.hanSuDung}}'),

        # Điện thoại email (thứ hai - chủ hộ mới)
        ('Điện thoại (nếu có): .......................................................  Email (nếu có): ......................',
         'Điện thoại: {{chuHoMoi.dienThoai}}  Email: {{chuHoMoi.email}}'),

        # Điện thoại và fax của hộ kinh doanh
        ('Điện thoại (nếu có): ..................................................................  Fax (nếu có): ..............',
         'Điện thoại: {{hoKinhDoanh.dienThoai}}  Fax: {{hoKinhDoanh.fax}}'),

        # Email và website của hộ kinh doanh
        ('Email (nếu có): ........................................................................  Website (nếu có): .......',
         'Email: {{hoKinhDoanh.email}}  Website: {{hoKinhDoanh.website}}'),
    ]

    replaced_count = 0
    found_patterns = []

    # Process paragraphs
    for para in doc.paragraphs:
        original = para.text
        modified = original

        for old, new in replacements:
            if old in modified:
                modified = modified.replace(old, new, 1)  # Replace only first occurrence
                found_patterns.append(old[:50])

        if modified != original:
            para.clear()
            para.add_run(modified)
            replaced_count += 1
            print(f"✓ [{replaced_count}] Replaced: {original[:60]}...")

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original = para.text
                    modified = original

                    for old, new in replacements:
                        if old in modified:
                            modified = modified.replace(old, new, 1)
                            found_patterns.append(old[:50])

                    if modified != original:
                        para.clear()
                        para.add_run(modified)
                        replaced_count += 1
                        print(f"✓ [{replaced_count}] Replaced in table")

    doc.save(output_file)
    print(f"\n✅ Total replacements: {replaced_count}")
    print(f"✅ Saved to: {output_file}")

    # Summary
    print(f"\n📋 Patterns found: {len(set(found_patterns))}")

if __name__ == "__main__":
    print("🔧 Adding COMPLETE placeholders to HKD_THAY_DOI.docx...\n")
    add_complete_placeholders()
